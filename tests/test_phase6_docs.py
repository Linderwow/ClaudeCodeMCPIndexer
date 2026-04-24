"""Phase 6: doc ingestion (MD + DOCX minimum; PDF tested via pypdf if available)."""
from __future__ import annotations

import asyncio
from pathlib import Path

import pytest

from code_rag.chunking.docs import DocChunker
from code_rag.config import load_settings
from code_rag.embedders.fake import FakeEmbedder
from code_rag.indexing.indexer import Indexer
from code_rag.stores.chroma_vector import ChromaVectorStore
from code_rag.stores.sqlite_lexical import SqliteLexicalStore

# ---- unit: DocChunker for markdown -----------------------------------------


def test_markdown_heading_aware_split(tmp_path: Path) -> None:
    p = tmp_path / "doc.md"
    p.write_text(
        """\
# Top

Intro paragraph.

## Inner A

Content A goes here with some text.

## Inner B

Content B goes here with some text.

### Inner B.1

Nested content.
""",
        encoding="utf-8",
    )
    ch = DocChunker(min_chars=5, max_chars=10_000)
    chunks = ch.chunk_file("repo", p, "doc.md")
    assert chunks, "should produce chunks"
    trails = {c.symbol for c in chunks}
    # Expect each heading to show up in the symbol breadcrumb.
    assert any(s and "Top" in s for s in trails)
    assert any(s and "Inner A" in s for s in trails)
    assert any(s and "Inner B" in s for s in trails)
    assert any(s and "Inner B.1" in s for s in trails)
    # All doc chunks are kind=doc.
    assert all(c.kind.value == "doc" for c in chunks)


def test_markdown_small_files_still_chunk(tmp_path: Path) -> None:
    p = tmp_path / "tiny.md"
    p.write_text("just a short note with enough content to clear min_chars", encoding="utf-8")
    ch = DocChunker(min_chars=10, max_chars=1000)
    chunks = ch.chunk_file("repo", p, "tiny.md")
    assert len(chunks) == 1
    assert chunks[0].language == "markdown"


# ---- integration: indexer picks up .md files --------------------------------


def test_indexer_picks_up_markdown(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    root = tmp_path / "repo"
    (root / "docs").mkdir(parents=True)
    (root / "docs" / "notes.md").write_text(
        "# Rare Heading XYZZY_MARKER\n\nbody text explaining the marker.\n",
        encoding="utf-8",
    )
    cfg = tmp_path / "config.toml"
    cfg.write_text(
        f"""
[paths]
roots    = ["{root.as_posix()}"]
data_dir = "{(tmp_path / 'data').as_posix()}"
log_dir  = "{(tmp_path / 'logs').as_posix()}"

[embedder]
kind     = "lm_studio"
base_url = "http://localhost:1234/v1"
model    = "qwen3-embedding-4b"

[reranker]
kind     = "lm_studio"
base_url = "http://localhost:1234/v1"
model    = "qwen3-reranker-4b"

[chunker]
min_chars = 5
max_chars = 2400
""",
        encoding="utf-8",
    )
    monkeypatch.setenv("CODE_RAG_CONFIG", str(cfg))
    settings = load_settings()

    embedder = FakeEmbedder(dim=32)
    vec = ChromaVectorStore(settings.chroma_dir, settings.vector_store.collection,
                            settings.index_meta_path)
    lex = SqliteLexicalStore(settings.fts_path)
    meta = ChromaVectorStore.build_meta("fake", embedder.model, embedder.dim)
    vec.open(meta)
    lex.open()
    try:
        indexer = Indexer(settings, embedder, vec, lexical_store=lex)
        stats = asyncio.run(indexer.reindex_all())
        assert stats.chunks_upserted > 0

        # Lexical hit for the rare marker confirms full pipeline (walker -> doc
        # chunker -> embed -> lex upsert).
        hits = lex.query("XYZZY_MARKER", k=5)
        assert hits
        assert hits[0].chunk.language == "markdown"
    finally:
        vec.close()
        lex.close()


# ---- unit: docx (via python-docx) ------------------------------------------


def test_docx_chunker_roundtrip(tmp_path: Path) -> None:
    # Build a minimal .docx on the fly.
    import docx  # python-docx
    p = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_paragraph("Alpha paragraph with the UNIQUE_DOCX_MARKER token.")
    d.add_paragraph("Beta paragraph.")
    d.save(str(p))
    ch = DocChunker(min_chars=5, max_chars=10_000)
    chunks = ch.chunk_file("repo", p, "sample.docx")
    assert chunks
    assert "UNIQUE_DOCX_MARKER" in chunks[0].text
    assert chunks[0].language == "docx"
